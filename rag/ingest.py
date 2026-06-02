"""
rag/ingest.py — Document loaders for ingestion into the knowledge base
=======================================================================
Extracts plain text from common document formats (PDF, Word, Markdown, plain
text, HTML) so PRDs / TDDs / specs can be ingested, chunked, and retrieved by
the agent. Heavy parsers (pypdf, python-docx) are imported lazily.
"""

from __future__ import annotations

from pathlib import Path

# Extensions we know how to extract text from.
SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".md", ".markdown", ".txt", ".rst", ".html", ".htm",
}


def is_supported(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_EXTENSIONS


def load_document(path: Path) -> str:
    """Extract text from a document. Raises ValueError for unsupported types."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)
    if suffix in (".html", ".htm"):
        from tools.web_tools import _html_to_text
        return _html_to_text(path.read_text(encoding="utf-8", errors="replace"), max_chars=10_000_000)
    if suffix in (".md", ".markdown", ".txt", ".rst"):
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(
        f"Unsupported document type '{suffix}'. Supported: "
        + ", ".join(sorted(SUPPORTED_EXTENSIONS))
    )


def _load_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:  # pragma: no cover
        raise ImportError("pypdf is not installed. Run: pip install pypdf") from e

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n\n".join(p for p in pages if p.strip())


def _load_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:  # pragma: no cover
        raise ImportError("python-docx is not installed. Run: pip install python-docx") from e

    document = docx.Document(str(path))
    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    # Include table cell text (PRDs often use tables)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
